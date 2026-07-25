from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path

workspace = Path(__file__).resolve().parent

# Build the Word document

doc = Document()
doc.add_heading('FizzBuzz App Deliverable', 1)
doc.add_paragraph('This document contains the FizzBuzz app source code and the expected output for the lab submission.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Source code file: ').bold = True
p.add_run('script.js')

code_text = '''const fizzBuzzList = document.getElementById("fizzBuzzList");

function generateFizzBuzz(limit = 20) {
    const results = [];

    for (let number = 1; number <= limit; number += 1) {
        if (number % 3 === 0 && number % 5 === 0) {
            results.push(`${number} FizzBuzz`);
        } else if (number % 3 === 0) {
            results.push(`${number} Fizz`);
        } else if (number % 5 === 0) {
            results.push(`${number} Buzz`);
        } else {
            results.push(String(number));
        }
    }

    return results;
}

function renderFizzBuzz() {
    const results = generateFizzBuzz();

    fizzBuzzList.innerHTML = "";
    results.forEach((result) => {
        const listItem = document.createElement("li");
        listItem.textContent = result;
        fizzBuzzList.appendChild(listItem);
    });

    console.clear();
    console.log("FizzBuzz Output:");
    results.forEach((result) => console.log(result));
}

renderFizzBuzz();'''

paragraph = doc.add_paragraph()
paragraph.add_run('Code:').bold = True
code_paragraph = doc.add_paragraph(code_text)
code_paragraph.style = 'Intense Quote'

doc.add_heading('Output Preview', 1)

# Create a simple screenshot-style image for the output preview
width, height = 800, 600
img = Image.new('RGB', (width, height), color=(248, 250, 252))
draw = ImageDraw.Draw(img)
try:
    font_title = ImageFont.load_default()
except TypeError:
    font_title = ImageFont.load_default()
font_body = ImageFont.load_default()

draw.text((40, 40), 'FizzBuzz Challenge', fill=(37, 99, 235))
draw.text((40, 80), 'Numbers 1 through 20', fill=(71, 85, 105))
lines = [
    '1',
    '2',
    '3 Fizz',
    '4',
    '5 Buzz',
    '6 Fizz',
    '7',
    '8',
    '9 Fizz',
    '10 Buzz',
    '11',
    '12 Fizz',
    '13',
    '14',
    '15 FizzBuzz',
    '16',
    '17',
    '18 Fizz',
    '19',
    '20 Buzz',
]
for index, line in enumerate(lines, start=1):
    draw.text((60, 130 + index * 20), line, fill=(15, 23, 42))

image_path = workspace / 'fizzbuzz_output.png'
img.save(image_path)
doc.add_picture(str(image_path), width=Inches(6.0))

output_path = workspace / 'FizzBuzz_Deliverable.docx'
doc.save(output_path)
print(output_path)
